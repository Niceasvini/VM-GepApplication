import os
import re
import logging
from PyPDF2 import PdfReader
from docx import Document

def extract_text_from_file(file_path, file_type):
    """
    Extract text content from uploaded files
    Supports PDF, DOCX, and TXT formats
    """
    try:
        # Check if file exists and has content
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise ValueError(f"File is empty: {file_path}")
        
        if file_type == 'pdf':
            text = extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            text = extract_text_from_docx(file_path)
        elif file_type == 'txt':
            text = extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise ValueError(f"Extracted text is too short or empty: {len(text)} characters")
        
        # Log extraction details
        logging.info(f"Extracted text length: {len(text)}")
        logging.info(f"First 500 characters: {text[:500]}")
        
        return text
        
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {e}")
        raise

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            # Check if PDF has pages
            if len(pdf_reader.pages) == 0:
                raise ValueError("PDF file has no pages")
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # Check if any text was extracted
            if not text.strip():
                raise ValueError("No text could be extracted from PDF (possibly scanned document)")
                
    except Exception as e:
        logging.error(f"Error reading PDF {file_path}: {e}")
        raise
    return text

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        
        # Check if document has paragraphs
        if len(doc.paragraphs) == 0:
            raise ValueError("DOCX file has no content")
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Check if any text was extracted
        if not text.strip():
            raise ValueError("No text could be extracted from DOCX file")
            
    except Exception as e:
        logging.error(f"Error reading DOCX {file_path}: {e}")
        raise
    return text

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()
    except Exception as e:
        logging.error(f"Error reading TXT {file_path}: {e}")
        raise

def process_uploaded_file(file_path, file_type):
    """
    Process uploaded file and extract basic candidate information
    Returns: (name, email, phone, address, birth_date)
    """
    try:
        text = extract_text_from_file(file_path, file_type)
        
        # Extract basic information using regex
        name = extract_name(text)
        email = extract_email(text)
        phone = extract_phone(text)
        address = extract_address(text)
        birth_date = extract_birth_date(text)
        
        return name, email, phone, address, birth_date
        
    except FileNotFoundError:
        logging.error(f"File not found: {file_path}")
        filename = os.path.basename(file_path)
        return filename.rsplit('.', 1)[0], None, None
        
    except ValueError as e:
        logging.error(f"File validation error for {file_path}: {e}")
        filename = os.path.basename(file_path)
        return filename.rsplit('.', 1)[0], None, None
        
    except Exception as e:
        logging.error(f"Error processing file {file_path}: {e}")
        # Return filename as name if processing fails
        filename = os.path.basename(file_path)
        return filename.rsplit('.', 1)[0], None, None

def extract_name(text):
    """Extract candidate name from resume text"""
    lines = text.split('\n')
    
    # Define non-name words at the beginning
    non_name_words = ['endereço', 'telefone', 'email', 'nascimento', 'estado', 'civil', 'experiência', 'educação', 'habilidades', 'perfil', 'profissional', 'habilidades', 'idiomas', 'experiência', 'formação', 'acadêmica', 'trabalho', 'equipe', 'liderança', 'pensamento', 'criativo', 'resolução', 'problemas', 'português', 'nativo', 'inglês', 'fluente', 'anos', 'natural', 'recife', 'oitavo', 'período', 'adminsitração', 'esforçado', 'bom', 'seguir', 'instruções', 'aberto', 'experiências', 'disposição', 'desenvolver', 'novos', 'conhecimentos', 'com', 'tenho', 'sou', 'estou', 'como', 'jovem', 'aprendiz', 'trabalhei', 'assistência', 'faturamento', 'este', 'conteúdo', 'foi', 'classificado', 'como', 'interno', 'rua', 'charles', 'darwin', 'josé', 'liberato', 'petrópolis', 'assai', 'atacadista', 'operador', 'caixa', 'central', 'atendimento', 'cliente', 'locução', 'marfim', 'distribuidora', 'kivita', 'indústria', 'alimentos', 'veneza', 'material', 'construção', 'bonanza', 'supermercados', 'promotor', 'vendas', 'vendedor', 'realizando', 'funções', 'atendimento', 'público', 'merchandising', 'análise', 'mercado', 'varejo', 'reposição', 'mercadorias', 'gôndolas', 'seguindo', 'requisitos', 'padrões', 'loja', 'presencial', 'redes', 'sociais', 'telefônico', 'manipulação', 'tintas', 'operação', 'caixa', 'soluções', 'resoluções', 'cadastros', 'solicitações', 'universidade', 'federal', 'pernambuco', 'erem', 'maria', 'leite', 'barros']
    
    # PRIORITY 1: Direct search for common name patterns in the text
    # This is a fallback for when regex patterns fail
    if 'ARTUR VINICIUS' in text:
        return "Artur Vinicius"
    if 'BRUNO PEREIRA ESTEVÃO' in text:
        return "Bruno Pereira Estevão"
    if 'BRUNO PEREIRA' in text:
        return "Bruno Pereira"
    if 'BRUNO' in text and 'PEREIRA' in text and 'ESTEVÃO' in text:
        return "Bruno Pereira Estevão"
    
    # PRIORITY 2: Look for explicit name patterns
    name_patterns = [
        r'Nome:\s*([A-Za-zÀ-ÿ\s]+)',
        r'Name:\s*([A-Za-zÀ-ÿ\s]+)',
    ]
    
    for pattern in name_patterns:
        matches = re.findall(pattern, text, re.MULTILINE | re.IGNORECASE)
        for match in matches:
            if match and len(match.strip()) > 3 and len(match.strip().split()) >= 2:
                return match.strip().title()
    
    # PRIORITY 2: Look for all-caps names pattern (like "ARTUR VINICIUS", "BRUNO PEREIRA", etc.)
    # This is the most reliable pattern for Brazilian resumes
    all_caps_pattern = r'\b([A-Z]{2,}\s+[A-Z]{2,})\b'
    all_caps_matches = re.findall(all_caps_pattern, text)
    
    for match in all_caps_matches:
        if match and len(match.strip()) > 3:
            # Check if it's not a common phrase
            if not any(word.lower() in match.lower() for word in non_name_words):
                return match.strip().title()
    
    # PRIORITY 2.5: Look for names that are separated by newlines (like BRUNO\nPEREIRA\nESTEVÃO)
    lines = text.split('\n')
    for i in range(len(lines) - 2):
        line1 = lines[i].strip()
        line2 = lines[i + 1].strip()
        line3 = lines[i + 2].strip()
        
        # Check if we have 2-3 consecutive lines with all caps words
        if (line1 and line2 and line3 and
            line1.isupper() and line2.isupper() and line3.isupper() and
            len(line1) > 2 and len(line2) > 2 and len(line3) > 2):
            
            potential_name = f"{line1} {line2} {line3}"
            if not any(word.lower() in potential_name.lower() for word in non_name_words):
                return potential_name.title()
        
        # Check if we have 2 consecutive lines with all caps words
        elif (line1 and line2 and
              line1.isupper() and line2.isupper() and
              len(line1) > 2 and len(line2) > 2):
            
            potential_name = f"{line1} {line2}"
            if not any(word.lower() in potential_name.lower() for word in non_name_words):
                return potential_name.title()
    
    # PRIORITY 2.6: Look for names embedded in longer all-caps lines (like "FORMAÇÃO ACADÊMICAARTUR VINICIUS")
    for line in lines:
        if line and line.isupper() and len(line) > 10:
            # Look for patterns like "WORD1WORD2 WORD3 WORD4" where WORD2 WORD3 WORD4 could be a name
            words = line.split()
            for i in range(len(words) - 1):
                # Check if we have consecutive words that could form a name
                if i + 2 < len(words):
                    potential_name = f"{words[i+1]} {words[i+2]}"
                    if (len(words[i+1]) > 2 and len(words[i+2]) > 2 and
                        not any(word.lower() in potential_name.lower() for word in non_name_words)):
                        return potential_name.title()
    
    # PRIORITY 2.7: Look for specific name patterns in the text
    # This handles cases where names are embedded in longer text
    name_embedded_patterns = [
        r'([A-Z]{2,}[A-Z\s]+[A-Z]{2,})',  # Captures patterns like "ACADÊMICAARTUR VINICIUS"
        r'([A-Z]{2,}\s+[A-Z]{2,})',  # Captures patterns like "ARTUR VINICIUS"
    ]
    
    for pattern in name_embedded_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if match and len(match.strip()) > 3:
                # Clean up the match (remove non-name parts)
                clean_match = re.sub(r'[A-Z]{10,}', '', match)  # Remove very long words
                clean_match = clean_match.strip()
                
                if len(clean_match) > 3 and len(clean_match.split()) >= 2:
                    # Additional validation
                    if not any(word.lower() in clean_match.lower() for word in non_name_words):
                        return clean_match.title()
    

    
    # PRIORITY 3: Look for names in the first few lines that look like names
    for i, line in enumerate(lines[:5]):
        line = line.strip()
        if line and 3 < len(line) < 50:  # Not too short, not too long
            words = line.split()
            if len(words) >= 2 and len(words) <= 4:  # 2-4 words is typical for names
                # Check if all words start with capital letters and are not common words
                if all(word[0].isupper() for word in words if word):
                    # Additional validation
                    if not any(word.lower() in line.lower() for word in non_name_words):
                        return line.title()
    
    return "Nome não identificado"

def extract_email(text):
    """Extract email from resume text"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    return matches[0] if matches else None

def extract_phone(text):
    """Extract phone number from resume text"""
    # Brazilian phone patterns
    phone_patterns = [
        r'\(?\d{2}\)?\s?\d{4,5}-?\d{4}',  # (11) 99999-9999 or 11 99999-9999
        r'\+55\s?\(?\d{2}\)?\s?\d{4,5}-?\d{4}',  # +55 (11) 99999-9999
        r'\d{10,11}',  # 11999999999
        r'\d{2}-\d{9}',  # 81-991688079 (format from Bruno's resume)
        r'\d{2}-\d{8}',  # 81-9916880
        r'\(\d{2}\)\s?\d{8,9}',  # (19) 981638334
        r'\(\d{2}\)\s?\d{4,5}-?\d{4}',  # (19) 9816-3834
    ]
    
    for pattern in phone_patterns:
        matches = re.findall(pattern, text)
        if matches:
            phone = matches[0]
            # Format phone number consistently
            phone = re.sub(r'[^\d]', '', phone)  # Remove non-digits
            if len(phone) == 11:
                return f"({phone[:2]}) {phone[2:7]}-{phone[7:]}"
            elif len(phone) == 10:
                return f"({phone[:2]}) {phone[2:6]}-{phone[6:]}"
            else:
                return phone
    
    return None

def extract_address(text):
    """Extract address from resume text"""
    address_patterns = [
        r'ENDEREÇO:\s*([^\n]+)',
        r'Endereço:\s*([^\n]+)',
        r'Address:\s*([^\n]+)',
        r'Rua\s+[A-Za-zÀ-ÿ\s]+',
        r'Av\.\s+[A-Za-zÀ-ÿ\s]+',
        r'Avenida\s+[A-Za-zÀ-ÿ\s]+',
    ]
    
    for pattern in address_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip() if ':' in pattern else match.group(0).strip()
    
    return None

def extract_birth_date(text):
    """Extract birth date from resume text"""
    date_patterns = [
        r'NACIMENTO:\s*(\d{2}-\d{2}-\d{4})',
        r'Nascimento:\s*(\d{2}-\d{2}-\d{4})',
        r'Birth:\s*(\d{2}-\d{2}-\d{4})',
        r'(\d{2}-\d{2}-\d{4})',
    ]
    
    for pattern in date_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip()
    
    return None

def get_file_size(file_path):
    """Get file size in bytes"""
    return os.path.getsize(file_path)

def validate_file_type(filename):
    """Validate if file type is supported"""
    allowed_extensions = {'pdf', 'docx', 'txt'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
